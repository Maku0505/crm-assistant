"""
ingest.py — Load, chunk, embed, and store all documents into ChromaDB.
Run once before starting the app:  python ingest.py

Install (no compiler required):
    pip install langchain langchain-community chromadb \
                sentence-transformers pypdf \
                python-docx openpyxl
"""

import email
import hashlib
import argparse
import zipfile
import tempfile
import shutil
from pathlib import Path

import docx                      # python-docx
import openpyxl
from pypdf import PdfReader

from langchain_core.documents import Document
from langchain_community.document_loaders import TextLoader, CSVLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_ollama import OllamaEmbeddings
from langchain_community.vectorstores import Chroma
from config import DATA_DIR, CHROMA_DIR, EMBED_MODEL, CHUNK_SIZE, CHUNK_OVERLAP


# ── Ingestion settings ────────────────────────────────────────────────────────

FOLDER_TYPE_MAP = {
    "crm_records": "crm",
    "sales":       "sales",
    "tickets":     "ticket",
    "emails":      "email",
    "documents":   "document",
    "documets":    "document",   # typo in dataset folder name
    "files":       "file",
}


# ── Lightweight loaders (no unstructured / llvmlite) ─────────────────────────

def load_txt(path: Path) -> list[Document]:
    return TextLoader(str(path), encoding="utf-8", autodetect_encoding=True).load()


def load_csv(path: Path) -> list[Document]:
    return CSVLoader(str(path)).load()


def load_pdf(path: Path) -> list[Document]:
    reader = PdfReader(str(path))
    docs = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            docs.append(Document(page_content=text, metadata={"page": i + 1}))
    return docs


def load_docx(path: Path) -> list[Document]:
    d = docx.Document(str(path))
    text = "\n\n".join(p.text for p in d.paragraphs if p.text.strip())
    return [Document(page_content=text, metadata={})] if text else []


def load_xlsx(path: Path) -> list[Document]:
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    docs = []
    for sheet in wb.worksheets:
        rows = []
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                rows.append("\t".join(cells))
        if rows:
            text = f"[Sheet: {sheet.title}]\n" + "\n".join(rows)
            docs.append(Document(page_content=text, metadata={}))
    return docs


def load_eml(path: Path) -> list[Document]:
    msg = email.message_from_bytes(path.read_bytes())
    headers = "\n".join(
        f"{h}: {msg.get(h, '')}"
        for h in ("From", "To", "Subject", "Date")
        if msg.get(h)
    )
    body = ""
    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_type() == "text/plain":
                body = part.get_payload(decode=True).decode(errors="replace")
                break
    else:
        body = msg.get_payload(decode=True).decode(errors="replace")
    text = (headers + "\n\n" + body).strip()
    return [Document(page_content=text, metadata={})] if text else []


def load_json(path: Path) -> list[Document]:
    import json
    try:
        data = json.loads(path.read_text(encoding="utf-8", errors="replace"))
    except Exception as e:
        print(f"  [WARN] Could not parse JSON {path.name}: {e}")
        return []
    docs = []
    # Handle list of email thread objects
    if isinstance(data, list):
        for item in data:
            if isinstance(item, dict):
                lines = []
                for k, v in item.items():
                    if isinstance(v, (str, int, float)):
                        lines.append(f"{k}: {v}")
                    elif isinstance(v, list):
                        # e.g. messages array inside a thread
                        for msg in v:
                            if isinstance(msg, dict):
                                lines.append("\n".join(f"{mk}: {mv}" for mk, mv in msg.items() if isinstance(mv, str)))
                text = "\n".join(lines).strip()
                if text:
                    docs.append(Document(page_content=text, metadata={}))
    elif isinstance(data, dict):
        text = "\n".join(f"{k}: {v}" for k, v in data.items() if isinstance(v, str))
        if text:
            docs.append(Document(page_content=text, metadata={}))
    return docs


# Extension → loader function
LOADER_MAP = {
    ".txt":  load_txt,
    ".md":   load_txt,
    ".csv":  load_csv,
    ".pdf":  load_pdf,
    ".docx": load_docx,
    ".doc":  load_docx,
    ".xlsx": load_xlsx,
    ".xls":  load_xlsx,
    ".eml":  load_eml,
    ".msg":  load_eml,
    ".json": load_json,
}


# ── Helpers ───────────────────────────────────────────────────────────────────

def detect_source_type(file_path: Path) -> str:
    parts = file_path.relative_to(DATA_DIR).parts
    return FOLDER_TYPE_MAP.get(parts[0], "general") if parts else "general"


def enrich_metadata(docs: list[Document], file_path: Path) -> list[Document]:
    source_type = detect_source_type(file_path)
    relative    = str(file_path.relative_to(DATA_DIR))
    for doc in docs:
        doc.metadata.update({
            "source_type":   source_type,
            "filename":      file_path.name,
            "relative_path": relative,
        })
    return docs


def load_file(file_path: Path) -> list[Document]:
    loader_fn = LOADER_MAP.get(file_path.suffix.lower())
    if loader_fn is None:
        return []
    try:
        return loader_fn(file_path)
    except Exception as e:
        print(f"  [WARN] Could not load {file_path.name}: {e}")
        return []


# ── Zip handling ──────────────────────────────────────────────────────────────

def unzip_archives(root: Path) -> int:
    """Recursively find all .zip files under root and extract them in-place.
    Each zip is extracted into a sibling folder named after the zip (without .zip).
    Already-extracted folders are skipped. Returns the number of zips extracted.
    """
    count = 0
    # Keep looping until no new zips are found (handles zips inside zips)
    while True:
        zips = [p for p in root.rglob("*.zip") if p.is_file()]
        new_this_round = 0
        for zp in zips:
            dest = zp.with_suffix("")          # e.g. data/emails/batch.zip → data/emails/batch
            if dest.exists():
                continue                        # already extracted, skip
            print(f"  [ZIP] Extracting {zp.relative_to(root)} → {dest.name}/")
            try:
                dest.mkdir(parents=True, exist_ok=True)
                with zipfile.ZipFile(zp, "r") as zf:
                    zf.extractall(dest)
                count += 1
                new_this_round += 1
            except zipfile.BadZipFile:
                print(f"  [WARN] Skipping bad zip: {zp.name}")
        if new_this_round == 0:
            break                              # no new zips discovered, done
    return count


# ── Main ──────────────────────────────────────────────────────────────────────

def main(reset: bool = False):
    if not DATA_DIR.exists():
        print(f"[ERROR] Data directory '{DATA_DIR}' not found.")
        return

    if reset and CHROMA_DIR.exists():
        shutil.rmtree(CHROMA_DIR)
        print("[INFO] Cleared existing ChromaDB.")

    # Unzip any archives before scanning for files
    n = unzip_archives(DATA_DIR)
    if n:
        print(f"[INFO] Extracted {n} zip archive(s).\n")

    all_files = [
        p for p in DATA_DIR.rglob("*")
        if p.is_file() and p.suffix.lower() in LOADER_MAP
    ]
    print(f"[INFO] Found {len(all_files)} file(s) to process.\n")

    all_docs = []
    for i, fp in enumerate(all_files, 1):
        print(f"  [{i}/{len(all_files)}] {fp.relative_to(DATA_DIR)}")
        docs = load_file(fp)
        docs = enrich_metadata(docs, fp)
        all_docs.extend(docs)

    if not all_docs:
        print("[ERROR] No documents loaded. Check your data folder.")
        return

    print(f"\n[INFO] Loaded {len(all_docs)} document(s). Storing whole files (no chunking)...")

    # Store each document as-is — no chunking.
    # This ensures every file is always retrieved completely.
    # For very large files (>8000 chars) we do a single split just to stay
    # within Mistral's context window, but most files will stay in one piece.
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

    chunks = []
    for doc in all_docs:
        if len(doc.page_content) <= CHUNK_SIZE:
            chunks.append(doc)   # fits whole — no split
        else:
            pieces = splitter.split_documents([doc])
            print(f"  [SPLIT] {doc.metadata.get('filename','?')} → {len(pieces)} pieces ({len(doc.page_content)} chars)")
            chunks.extend(pieces)

    print(f"[INFO] {len(chunks)} chunk(s) total ({len(all_docs)} files).")

    ids = [
        hashlib.md5(
            f"{c.metadata.get('relative_path','?')}:{i}:{c.page_content[:80]}".encode()
        ).hexdigest()
        for i, c in enumerate(chunks)
    ]

    print(f"\n[INFO] Loading embedding model via Ollama (make sure 'ollama serve' is running)...")
    embeddings = OllamaEmbeddings(
        model="nomic-embed-text",
    )

    print("[INFO] Building ChromaDB (this may take a few minutes)...")
    Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        ids=ids,
        persist_directory=str(CHROMA_DIR),
    )

    print(f"\n[OK] Done. {len(chunks)} chunks stored in '{CHROMA_DIR}'.")
    print("     Run:  streamlit run app.py")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--reset", action="store_true",
                        help="Wipe ChromaDB and rebuild from scratch.")
    main(reset=parser.parse_args().reset)